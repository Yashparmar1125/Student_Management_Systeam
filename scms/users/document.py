from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()
doc.add_heading('Phase 1 Study Plan - Full Stack Development & DSA', 0)

# Add Full Stack Development Timetable
doc.add_heading('Full Stack Development Study Plan', level=1)

# Table for Full Stack Development Study Plan
full_stack_table = [
    ["Week", "Day", "Topic", "Subtopics", "Project"],
    ["Week 1", "Day 1-2", "HTML, CSS Basics", "- Basic Structure, Tags, Attributes\n- Flexbox, Grid", "Build a personal website"],
    ["", "Day 3-4", "Advanced CSS", "- CSS Animations, Transitions\n- CSS Variables, Responsive Design", "Enhance personal website design"],
    ["", "Day 5-6", "JavaScript Basics", "- Variables, Loops, Functions\n- DOM Manipulation, Events", "Build a dynamic portfolio website"],
    ["Week 2", "Day 7-8", "JavaScript Advanced", "- Closures, Promises, Async/Await\n- Error Handling", "Portfolio with API Fetching"],
    ["", "Day 9-10", "Version Control (Git)", "- Git Basics, Branching, Merging\n- GitHub and GitLab", "Deploy your portfolio on GitHub"],
    ["", "Day 11-12", "Node.js Basics", "- Introduction to Node.js\n- Modules, NPM", "Create a basic Node.js server"],
    ["", "Day 13-14", "Express.js", "- Routing, Middleware\n- Error Handling, REST APIs", "Develop an API for your portfolio"],
    ["Week 3", "Day 15-16", "Database (MongoDB)", "- NoSQL, MongoDB Basics\n- CRUD Operations", "Connect MongoDB to your API"],
    ["", "Day 17-18", "Authentication", "- JWT, Passport.js\n- Session Management", "Add login and registration functionality"],
    ["", "Day 19-20", "Advanced JavaScript (ES6+) and Node.js", "- Async/Await, Destructuring\n- REST APIs and Middleware", "Enhance your portfolio app with more features"],
    ["Week 4", "Day 21-22", "React.js Basics", "- JSX, Components\n- Props, State, Lifecycle", "Build a basic React app"],
    ["", "Day 23-24", "React.js Advanced", "- Hooks, Context API, Redux\n- Component Design Patterns", "Enhance React app functionality"],
    ["", "Day 25-26", "Backend Integration (Node.js and React)", "- Connecting React to API\n- Authentication in React", "Connect your React app with backend"],
    ["Week 5", "Day 27-28", "Deployment", "- Deployment with Heroku, Netlify\n- CI/CD Pipelines", "Deploy your full stack application"],
    ["", "Day 29-30", "Final Project", "- Build a complete full-stack app", "Build a final full-stack project (e.g., E-commerce website)"],
    ["Week 6", "Day 31-33", "Project Review & Optimization", "- Review code, performance optimization\n- Clean up code and documentation", "Final project optimization and review"],
]

# Adding the table for Full Stack Development Study Plan
full_stack_table_columns = full_stack_table[0]
table = doc.add_table(rows=1, cols=len(full_stack_table_columns))

# Add column headers
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(full_stack_table_columns):
    hdr_cells[i].text = column_name

# Add data rows
for row in full_stack_table[1:]:
    row_cells = table.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = str(cell_value)

# Apply the borders to the table (correct method for setting borders)
tbl = table._element
tblBorders = OxmlElement('w:tblBorders')
for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border_element = OxmlElement(f'w:{border}')
    border_element.set('w:val', 'single')  # Set border style to 'single'
    border_element.set('w:sz', '4')  # Set border size to 4 (can adjust)
    border_element.set('w:space', '0')  # Set space between cells
    tblBorders.append(border_element)

tbl.tblBorders = tblBorders

# Add DSA Timetable
doc.add_page_break()
doc.add_heading('DSA Study Plan', level=1)

# Table for DSA Study Plan
dsa_table = [
    ["Week", "Day", "Topic", "Subtopics", "Project"],
    ["Week 1", "Day 1-2", "Introduction to DSA", "- What is DSA\n- Time Complexity\n- Space Complexity", "Solve basic problems on Array & String"],
    ["", "Day 3-4", "Arrays", "- Linear Search, Binary Search\n- Sorting Algorithms (Bubble, Selection, Insertion)", "Solve problems on Array Manipulation, Sorting Algorithms"],
    ["", "Day 5-6", "Arrays: Advanced Concepts", "- Sliding Window, Kadane’s Algorithm\n- Searching in Rotated Sorted Arrays", "Solve problems on Sliding Window, Kadane's Algorithm"],
    ["Week 2", "Day 7-8", "String Manipulation", "- Pattern Matching, Anagram Check\n- Palindrome Check", "Solve problems on Pattern Matching, Palindrome"],
    ["", "Day 9-10", "Strings: Advanced Algorithms", "- KMP Algorithm, Rabin-Karp Algorithm\n- Z-Algorithm, Trie", "Solve problems on KMP, Rabin-Karp"],
    ["", "Day 11-12", "Linked Lists", "- Linked List Basics, Singly Linked List\n- Doubly Linked List, Circular Linked List", "Solve problems on Reversal of Linked List, Detecting Loops"],
    ["", "Day 13-14", "Linked List: Advanced Problems", "- Merge Sort on Linked List\n- Floyd's Cycle-Finding Algorithm", "Solve problems on Intersection of Linked Lists, Merge Sort"],
    ["Week 3", "Day 15-16", "Stacks", "- Stack Basics, Stack Operations\n- Infix to Postfix and Prefix Conversion", "Solve problems on Infix to Postfix, Parenthesis Matching"],
    ["", "Day 17-18", "Queues", "- Circular Queue, Priority Queue\n- Deque", "Solve problems on Queue Implementations, Circular Queue"],
    ["", "Day 19-20", "Recursion", "- Basic Recursion, Backtracking\n- Divide and Conquer", "Solve problems on N-Queens Problem, Tower of Hanoi"],
    ["Week 4", "Day 21-22", "Trees: Basics", "- Binary Trees, Tree Traversals\n- Pre-order, In-order, Post-order", "Solve problems on Tree Traversals, Height of Tree"],
    ["", "Day 23-24", "Binary Search Tree (BST)", "- BST Operations, Search, Insert, Delete\n- Balanced Trees (AVL, Red-Black)", "Solve problems on Insertion/Deletion in BST"],
    ["", "Day 25-26", "Binary Search Tree (BST) Advanced", "- LCA (Lowest Common Ancestor)\n- Successor/Predecessor in BST", "Solve problems on LCA, BST Successor/Predecessor"],
    ["Week 5", "Day 27-28", "Heap", "- Min-Heap, Max-Heap\n- Heapify, Priority Queue", "Solve problems on Heap Construction, Heap Sort"],
    ["", "Day 29-30", "Graphs", "- BFS, DFS\n- Graph Representation", "Solve problems on BFS, DFS"],
    ["Week 6", "Day 31-32", "Graphs: Advanced Concepts", "- Dijkstra, Bellman-Ford\n- Topological Sort, SCC", "Solve problems on Dijkstra’s Algorithm, Topological Sort"],
    ["", "Day 33-34", "Dynamic Programming (DP): Basics", "- Recursion vs Memoization\n- Fibonacci Series, 0/1 Knapsack", "Solve problems on Fibonacci, Knapsack Problem"],
    ["", "Day 35-36", "Dynamic Programming (DP): Intermediate", "- LCS, LIS\n- Matrix Chain Multiplication", "Solve problems on LCS, LIS, Matrix Chain Multiplication"],
    ["", "Day 37-38", "Segment Tree", "- Tree Construction, Range Queries\n- Lazy Propagation", "Solve problems on Segment Tree, Union-Find"],
    ["", "Day 45-46", "Final Review & Project", "- Review all concepts, Mock Interviews\n- Solve a mix of DSA problems", "Complete final DSA project and optimize code"],
]

# Adding the table for DSA Study Plan
dsa_table_columns = dsa_table[0]
table = doc.add_table(rows=1, cols=len(dsa_table_columns))

# Add column headers
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(dsa_table_columns):
    hdr_cells[i].text = column_name

# Add data rows
for row in dsa_table[1:]:
    row_cells = table.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = str(cell_value)

# Apply the borders to the DSA table
tbl = table._element
tblBorders = OxmlElement('w:tblBorders')
for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border_element = OxmlElement(f'w:{border}')
    border_element.set('w:val', 'single')
    border_element.set('w:sz', '4')  # Set border size to 4
    border_element.set('w:space', '0')
    tblBorders.append(border_element)

tbl.tblBorders = tblBorders

# Save the document
doc.save('Study_Plan_Full_Stack_and_DSA.docx')
